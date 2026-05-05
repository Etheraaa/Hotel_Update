from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


OUTPUT_PATH = "/Users/bytedance/Documents/New project/VB1-run/酒店升房系统产品文档.docx"
PROJECT_PATH = "/Users/bytedance/Documents/New project/VB1-run"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "C8CDD4")


def add_paragraph(document: Document, text: str, bold=False, size=10.5, space_after=4):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    return paragraph


def add_title(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.alignment = 0
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(18)


def add_one_cell_table(document: Document, text: str):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = True
    set_cell_text(table.cell(0, 0), text)
    set_table_borders(table)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_grid_table(document: Document, headers, rows, column_widths_cm):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(column_widths_cm[idx])
        set_cell_text(cell, header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Cm(column_widths_cm[idx])
            set_cell_text(cells[idx], value)

    document.add_paragraph().paragraph_format.space_after = Pt(2)


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(10.5)

    add_title(document, "酒店升房系统项目产品文档")
    add_paragraph(document, "项目名称：酒店升房情报网站 V1", bold=True)
    add_paragraph(document, "项目定位：围绕高端酒店升房信息查询、入住观察投稿与升房沟通辅助的轻量情报产品。")
    add_paragraph(document, "参考基础：本文按照参考文件的成文方式整理，保留“说明文字 + 表格 + 启动信息”的文档结构。")

    add_paragraph(document, "项目路径：", bold=True)
    add_one_cell_table(document, PROJECT_PATH)
    add_paragraph(document, "本地启动说明：", bold=True)
    add_one_cell_table(document, "前端：cd frontend 之后执行 npm install 与 npm run dev")
    add_one_cell_table(document, "后端：cd backend 之后执行 npm install、npm run db:reset 与 npm run dev")
    add_one_cell_table(document, "访问地址：前端 http://localhost:3000 ，后端 http://localhost:4000")

    add_paragraph(document, "1. 产品定位与业务目标", bold=True, size=12, space_after=6)
    add_paragraph(document, "1.1. 产品定位与核心价值", bold=True)
    add_paragraph(document, "“酒店升房情报网站”是一款面向高频住店用户、酒店会员和出行决策人群的升房情报工具。")
    add_paragraph(document, "其核心价值在于：")
    add_paragraph(document, "降低决策不确定性：把分散在论坛、社群和个人经验里的升房观察整理成可筛选、可比较的酒店样本。")
    add_paragraph(document, "提升入住预期管理能力：让用户在预订前就能了解不同酒店、不同会员等级的大致升级方向。")
    add_paragraph(document, "沉淀可复用经验：通过投稿机制持续补充真实入住观察，形成可迭代的升房情报池。")
    add_paragraph(document, "辅助沟通表达：通过话术生成功能，把用户诉求整理成更自然、礼貌且更容易落地的表达。")

    add_paragraph(document, "1.2. 目标用户与核心痛点", bold=True)
    add_paragraph(document, "目标用户：")
    add_paragraph(document, "高频入住国际连锁酒店、关注会员权益兑现和房型升级机会的差旅及休闲住客。")
    add_paragraph(document, "核心痛点：")
    add_paragraph(document, "同一家酒店在不同会员等级、日期和房态下，升级结果差异较大，用户缺少系统化参考。")
    add_paragraph(document, "现有经验大多散落在论坛或聊天记录中，检索成本高，也缺少统一口径。")
    add_paragraph(document, "用户知道自己想争取更好房型或景观，但不一定能快速整理出得体的话术。")

    add_paragraph(document, "2. 产品范围与用户旅程", bold=True, size=12, space_after=6)
    add_paragraph(document, "2.1. 本期功能范围", bold=True)
    add_paragraph(document, "当前版本已经形成“查信息 - 看详情 - 补样本 - 生成沟通话术”的核心闭环。")
    add_grid_table(
        document,
        ["页面名称", "功能", "访问入口"],
        [
            ["首页", "提供酒店列表、关键词搜索，以及集团 / 品牌 / 城市三类筛选能力；展示每家酒店样本量和编辑摘要。", "/"],
            ["酒店详情页", "展示单酒店基础信息、来源说明、编辑备注，以及按会员等级拆分的升房分布表。", "/hotel/[hotelId]"],
            ["投稿页", "支持搜索酒店、选择会员等级、选择预订房型与最终房型、填写入住日期和情境，并提交真实观察。", "/submit"],
            ["话术生成页", "支持选择酒店、入住情境、会员等级、目标诉求和语气，生成一段可直接发送给酒店的话术。", "/generator"],
            ["接口层", "提供酒店查询、筛选、详情、升房统计、投稿元数据、房型联动和话术生成接口。", "/api/*"],
        ],
        [3.0, 9.8, 3.2],
    )

    add_paragraph(document, "2.2. 核心用户旅程", bold=True)
    add_paragraph(document, "用户首先在首页按酒店名、集团、品牌或城市筛选目标酒店，并快速浏览样本量与摘要判断是否值得深入查看。")
    add_paragraph(document, "进入酒店详情页后，用户查看来源说明、编辑备注和各会员等级升房分布，从而判断自己的会员等级大致可能对应什么升级方向。")
    add_paragraph(document, "若用户有新的入住观察，可以进入投稿页补充酒店、房型、日期与情境，帮助系统持续更新样本池。")
    add_paragraph(document, "若用户在入住前需要沟通辅助，则进入话术生成页，将“会员等级 + 诉求 + 情境”整理成一句可直接使用的话。")

    add_paragraph(document, "3. 功能设计与业务规则", bold=True, size=12, space_after=6)
    add_paragraph(document, "3.1. 核心功能设计要点", bold=True)
    add_paragraph(document, "3.1.1. 首页", bold=True)
    add_paragraph(document, "信息维度：每条酒店卡片包含城市、集团、酒店名称、样本量和一段摘要判断。")
    add_paragraph(document, "核心交互：支持关键词搜索，支持集团 / 品牌 / 城市筛选，结果为空时展示空状态提示。")
    add_paragraph(document, "排序逻辑：酒店列表按样本量从高到低优先展示，再按名称排序，以保证高参考价值酒店优先被看到。")

    add_paragraph(document, "3.1.2. 酒店详情页", bold=True)
    add_paragraph(document, "基础信息：展示酒店所属集团、品牌、城市、最新观察日期、来源池说明和编辑备注。")
    add_paragraph(document, "核心表格：按会员等级纵向展开，按“尊贵豪华房 / 小型套房 / 更高套房 / 特色房型”横向展示。")
    add_paragraph(document, "单元格口径：每个单元格展示“样本数 + 在该会员等级成功升级样本中的占比”，便于用户同时判断绝对量和相对倾向。")
    add_paragraph(document, "兜底机制：若当前酒店暂无足够统计样本，则展示“当前样本不足”的空状态，而不是输出误导性结论。")

    add_paragraph(document, "3.1.3. 投稿页", bold=True)
    add_paragraph(document, "填写流程：先搜索并选定酒店，再联动加载该酒店可选房型，之后补充会员等级、预订房型、最终房型、入住日期和入住情境。")
    add_paragraph(document, "体验目标：整页只保留必要字段，尽量让用户在 1 分钟内完成一条入住观察的补充。")
    add_paragraph(document, "成功反馈：提交成功后即时提示“投稿成功，感谢补充这条入住观察”，并重置表单进入下一次录入状态。")

    add_paragraph(document, "3.1.4. 话术生成页", bold=True)
    add_paragraph(document, "输入条件：酒店、入住情境、会员等级、目标诉求、沟通语气，以及可选的补充背景。")
    add_paragraph(document, "当前预设：入住情境包括纪念日、首次入住、晚到、只住一晚、想要安静、更好景观；目标诉求包括房型升级、更高楼层、更好景观；语气包括礼貌自然、稍微主动、商务正式。")
    add_paragraph(document, "输出结果：返回一段可直接复制的酒店沟通话术，优先强调自然礼貌和表达体面。")
    add_paragraph(document, "异常处理：若模型服务尚未配置完成，页面会给出明确提示，而不是返回空白结果。")

    add_paragraph(document, "3.2. 关键业务规则", bold=True)
    add_paragraph(document, "酒店与房型必须来自系统已有数据：投稿和话术功能都依赖已存在的酒店库，避免自由输入导致统计口径混乱。")
    add_paragraph(document, "只有可归类房型才计入升级统计：预订房型可以是基础房，但最终房型必须能映射到既定 room bucket，才能进入聚合计算。")
    add_paragraph(document, "入住日期必须为有效日期：系统校验 observed_at 字段格式，防止脏数据进入统计池。")
    add_paragraph(document, "投稿会即时反哺统计：一条投稿成功后，会更新酒店最新观察日期、样本汇总，以及该会员等级下的房型升级占比。")
    add_paragraph(document, "话术生成只在合法组合下执行：酒店、会员等级、情境、诉求和语气都必须命中预设范围，保证输出内容稳定可控。")

    add_paragraph(document, "4. 技术实现路径（通过 Next.js + Express + SQLite 实现）", bold=True, size=12, space_after=6)
    add_paragraph(document, "4.1. 整体架构", bold=True)
    add_paragraph(document, "前端技术栈：采用 Next.js App Router 构建前台站点，页面包括首页、酒店详情页、投稿页和话术生成页。")
    add_paragraph(document, "后端技术栈：采用 Express 提供 REST API 服务，统一处理酒店查询、投稿和话术生成请求。")
    add_paragraph(document, "数据存储：当前使用 SQLite，本地种子数据覆盖 8 家酒店，跨越万豪、希尔顿、凯悦、半岛等多个集团。")

    add_paragraph(document, "4.2. 核心接口", bold=True)
    add_paragraph(document, "酒店查询接口：GET /api/hotels、GET /api/hotels/filters、GET /api/hotels/:hotelId、GET /api/hotels/:hotelId/upgrade-stats。")
    add_paragraph(document, "投稿接口：GET /api/submissions/meta、GET /api/submissions/hotels、GET /api/submissions/hotels/:hotelId/room-options、POST /api/submissions。")
    add_paragraph(document, "话术接口：GET /api/phrasing/meta、POST /api/phrasing/generate。")

    add_paragraph(document, "4.3. AI 能力接入", bold=True)
    add_paragraph(document, "当前 AI 能力主要用于升房沟通话术生成，后端通过 OpenAI-compatible 的 responses client 发起文本生成请求。")
    add_paragraph(document, "Prompt 结构围绕酒店名称、会员等级、用户诉求、入住情境和补充背景组织，目标是生成简短、自然、不过度用力的表达。")
    add_paragraph(document, "若上游模型未配置或返回空内容，系统会显式报错并在前端给出可理解的失败提示。")

    add_paragraph(document, "4.4. 核心数据实体", bold=True)
    add_grid_table(
        document,
        ["实体名称", "字段列表", "说明"],
        [
            ["酒店实体 (Hotel)", "hotel_id, hotel_name, hotel_group, hotel_brand, city, sample_count, latest_observed_at, summary_text", "定义酒店的展示信息、样本规模和编辑摘要。"],
            ["升房统计实体 (Upgrade Stats)", "member_tier, room_bucket, success_count, success_ratio, tier_success_total", "定义单酒店在不同会员等级下的升级流向与占比。"],
            ["投稿实体 (Submission)", "hotel_id, member_tier, booked_room_raw, upgraded_room_raw, observed_at, stay_context", "记录单条真实入住观察，并作为统计更新来源。"],
            ["话术请求实体 (Phrasing Payload)", "hotel_id, scenario_ids, membership_level, goal_request, tone, additional_context", "定义生成升房沟通话术时需要的输入条件。"],
        ],
        [3.4, 7.0, 5.6],
    )

    add_paragraph(document, "4.5. 当前版本边界", bold=True)
    add_paragraph(document, "当前样本仍以种子数据和审核投稿为主，适合做“方向性判断”，暂不适合作为精确承诺工具。")
    add_paragraph(document, "项目暂未提供后台运营界面、投稿审核工作台和自动抓取能力，数据维护仍依赖手工种子更新与用户投稿。")
    add_paragraph(document, "话术生成是辅助功能，不直接承诺升房结果，其价值在于帮助用户把诉求说得更自然。")

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
